from flask import Flask, render_template, request, send_file
import requests
from bs4 import BeautifulSoup
import pandas as pd
from fpdf import FPDF
from docx import Document
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


app = Flask(__name__)

def scrape_data(url):
    response = requests.get(url)
    if response.status_code == 200:
        page_content = response.text
    else:
        return None, f"Failed to retrieve the web page. Status code: {response.status_code}"

    soup = BeautifulSoup(page_content, 'html.parser')

    headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    heading_texts = [heading.text.strip() for heading in headings]

    paragraphs = soup.find_all('p')
    paragraph_texts = [paragraph.text.strip() for paragraph in paragraphs]

    list_items = soup.find_all('li')
    list_item_texts = [item.text.strip() for item in list_items]

    return {
        'headings': heading_texts,
        'paragraphs': paragraph_texts,
        'list_items': list_item_texts
    }, None

def create_csv(data):
    df = pd.DataFrame({
        'Headings': data['headings'],
        'Paragraphs': data['paragraphs'],
        'List Items': data['list_items']
    })
    
    buffer = io.StringIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)
    
    return buffer.getvalue()


def create_pdf(data):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica", 12)
    y = height - 40

    # Add Headings
    c.drawString(40, y, "Headings:")
    y -= 20
    for heading in data['headings']:
        c.drawString(40, y, heading)
        y -= 15
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 40

    # Add Paragraphs
    c.drawString(40, y, "Paragraphs:")
    y -= 20
    for paragraph in data['paragraphs']:
        c.drawString(40, y, paragraph)
        y -= 15
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 40

    # Add List Items
    c.drawString(40, y, "List Items:")
    y -= 20
    for item in data['list_items']:
        c.drawString(40, y, item)
        y -= 15
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 40

    c.save()
    buffer.seek(0)
    return buffer

def create_docx(data):
    doc = Document()
    doc.add_heading('Headings', level=1)
    for heading in data['headings']:
        doc.add_paragraph(heading)
    
    doc.add_heading('Paragraphs', level=1)
    for paragraph in data['paragraphs']:
        doc.add_paragraph(paragraph)
    
    doc.add_heading('List Items', level=1)
    for item in data['list_items']:
        doc.add_paragraph(item)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/', methods=['GET', 'POST'])
def index():
    data = ""
    if request.method == 'POST':
        url = request.form.get('url')
        if url:
            scraped_data, error = scrape_data(url)
            if error:
                return render_template('index.html', data=error)
            
            data = scraped_data
            return render_template('index.html', data=data)

    return render_template('index.html', data=data)

@app.route('/download_csv', methods=['POST'])
def download_csv():
    csv_data = create_csv(scraped_data)
    return send_file(io.BytesIO(csv_data.encode()), mimetype='text/csv', as_attachment=True, download_name='data.csv')

@app.route('/download/pdf', methods=['POST'])
def download_pdf():
    url = request.form.get('url')
    scraped_data, error = scrape_data(url)
    if error:
        return error, 400

    pdf_data = create_pdf(scraped_data)
    return send_file(
        pdf_data,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='scraped_data.pdf'
    )

@app.route('/download/docx', methods=['POST'])
def download_docx():
    url = request.form.get('url')
    scraped_data, error = scrape_data(url)
    if error:
        return error, 400

    docx_data = create_docx(scraped_data)
    return send_file(
        docx_data,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='scraped_data.docx'
    )

if __name__ == '__main__':
    app.run(debug=True)
